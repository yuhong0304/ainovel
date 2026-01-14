"""
导出功能模块
支持导出为 TXT, DOCX, EPUB 格式
"""

import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Optional, Any
from datetime import datetime

logger = logging.getLogger(__name__)


class NovelExporter:
    """
    小说导出器
    
    支持格式:
    - TXT: 纯文本格式
    - DOCX: Word 文档格式
    - EPUB: 电子书格式
    """
    
    def __init__(self, project_path: Path):
        """
        初始化导出器
        
        Args:
            project_path: 项目目录路径
        """
        self.project_path = Path(project_path)
        self.content_dir = self.project_path / "content"
        self.export_dir = self.project_path / "exports"
        self.export_dir.mkdir(exist_ok=True)
        
        # 加载项目元数据
        self.meta = self._load_meta()
        
    def _load_meta(self) -> Dict[str, Any]:
        """加载项目元数据"""
        meta_file = self.project_path / "meta.json"
        if meta_file.exists():
            import json
            with open(meta_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {"title": self.project_path.name, "author": "佚名"}
    
    def _get_chapters(self) -> List[Dict[str, Any]]:
        """获取所有章节内容，按顺序排列"""
        chapters = []
        
        if not self.content_dir.exists():
            return chapters
            
        # 查找所有章节文件
        chapter_files = sorted(
            self.content_dir.glob("chapter_*.md"),
            key=lambda x: self._extract_chapter_num(x.name)
        )
        
        for cf in chapter_files:
            content = cf.read_text(encoding='utf-8')
            # 提取标题（如果有）
            title_match = re.match(r'^#\s*(.+)$', content, re.MULTILINE)
            title = title_match.group(1) if title_match else f"第{self._extract_chapter_num(cf.name)}章"
            
            # 清理内容（移除标题行）
            if title_match:
                content = content[title_match.end():].strip()
            
            chapters.append({
                "num": self._extract_chapter_num(cf.name),
                "title": title,
                "content": content,
                "file": cf.name
            })
            
        return chapters
    
    def _extract_chapter_num(self, filename: str) -> int:
        """从文件名提取章节号"""
        match = re.search(r'chapter_(\d+)', filename)
        return int(match.group(1)) if match else 0
    
    def export_txt(self, filename: Optional[str] = None) -> Path:
        """
        导出为 TXT 格式
        
        Args:
            filename: 可选的文件名
            
        Returns:
            导出文件路径
        """
        chapters = self._get_chapters()
        title = self.meta.get("title", "未命名小说")
        author = self.meta.get("author", "佚名")
        
        if not filename:
            filename = f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        
        output_path = self.export_dir / filename
        
        with open(output_path, 'w', encoding='utf-8') as f:
            # 写入标题和作者
            f.write(f"{title}\n")
            f.write(f"作者：{author}\n")
            f.write("=" * 50 + "\n\n")
            
            # 写入各章节
            for ch in chapters:
                f.write(f"\n{ch['title']}\n")
                f.write("-" * 30 + "\n\n")
                f.write(ch['content'])
                f.write("\n\n")
        
        logger.info(f"已导出 TXT: {output_path}")
        return output_path
    
    def export_docx(self, filename: Optional[str] = None) -> Path:
        """
        导出为 DOCX 格式
        
        需要安装: pip install python-docx
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("请先安装 python-docx: pip install python-docx")
        
        chapters = self._get_chapters()
        title = self.meta.get("title", "未命名小说")
        author = self.meta.get("author", "佚名")
        
        if not filename:
            filename = f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        output_path = self.export_dir / filename
        
        doc = Document()
        
        # 添加标题
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加作者
        author_para = doc.add_paragraph(f"作者：{author}")
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        # 添加各章节
        for ch in chapters:
            # 章节标题
            doc.add_heading(ch['title'], level=1)
            
            # 章节内容（按段落分割）
            paragraphs = ch['content'].split('\n\n')
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    para = doc.add_paragraph()
                    para.paragraph_format.first_line_indent = Inches(0.3)
                    run = para.add_run(para_text)
                    run.font.size = Pt(12)
            
            # 章节间分页
            doc.add_page_break()
        
        doc.save(output_path)
        logger.info(f"已导出 DOCX: {output_path}")
        return output_path
    
    def export_epub(self, filename: Optional[str] = None) -> Path:
        """
        导出为 EPUB 格式
        
        需要安装: pip install ebooklib
        """
        try:
            from ebooklib import epub
        except ImportError:
            raise ImportError("请先安装 ebooklib: pip install ebooklib")
        
        chapters = self._get_chapters()
        title = self.meta.get("title", "未命名小说")
        author = self.meta.get("author", "佚名")
        
        if not filename:
            filename = f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.epub"
        
        output_path = self.export_dir / filename
        
        # 创建 EPUB 书籍
        book = epub.EpubBook()
        book.set_identifier(f"novel-agent-{datetime.now().timestamp()}")
        book.set_title(title)
        book.set_language('zh-CN')
        book.add_author(author)
        
        # 创建章节
        epub_chapters = []
        for ch in chapters:
            # 创建章节对象
            epub_ch = epub.EpubHtml(
                title=ch['title'],
                file_name=f"chapter_{ch['num']}.xhtml",
                lang='zh-CN'
            )
            
            # 转换内容为 HTML
            content_html = self._markdown_to_html(ch['content'])
            epub_ch.content = f'''
            <html>
            <head><title>{ch['title']}</title></head>
            <body>
                <h1>{ch['title']}</h1>
                {content_html}
            </body>
            </html>
            '''
            
            book.add_item(epub_ch)
            epub_chapters.append(epub_ch)
        
        # 创建目录
        book.toc = [(epub.Section('目录'), epub_chapters)]
        
        # 添加导航文件
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # 定义阅读顺序
        book.spine = ['nav'] + epub_chapters
        
        # 添加默认样式
        style = '''
        body { font-family: "宋体", SimSun, serif; line-height: 1.8; }
        h1 { text-align: center; margin-bottom: 2em; }
        p { text-indent: 2em; margin: 0.5em 0; }
        '''
        nav_css = epub.EpubItem(
            uid="style_nav",
            file_name="style/nav.css",
            media_type="text/css",
            content=style
        )
        book.add_item(nav_css)
        
        # 写入文件
        epub.write_epub(output_path, book)
        logger.info(f"已导出 EPUB: {output_path}")
        return output_path
    
    def _markdown_to_html(self, markdown_text: str) -> str:
        """Markdown 转 HTML (使用 markdown 库)"""
        try:
            import markdown
            return markdown.markdown(markdown_text)
        except ImportError:
            # Fallback if library missing
            logger.warning("Markdown library not found, using simple fallback.")
            html_parts = []
            paragraphs = markdown_text.split('\n\n')
            for para in paragraphs:
                para = para.strip().replace('\n', '<br/>')
                if para:
                    html_parts.append(f'<p>{para}</p>')
            return '\n'.join(html_parts)
    
    def export_all(self) -> Dict[str, Path]:
        """导出所有格式"""
        results = {}
        
        try:
            results['txt'] = self.export_txt()
        except Exception as e:
            logger.error(f"TXT 导出失败: {e}")
            
        try:
            results['docx'] = self.export_docx()
        except Exception as e:
            logger.warning(f"DOCX 导出失败 (可能需要安装 python-docx): {e}")
            
        try:
            results['epub'] = self.export_epub()
        except Exception as e:
            logger.warning(f"EPUB 导出失败 (可能需要安装 ebooklib): {e}")
            
        return results
